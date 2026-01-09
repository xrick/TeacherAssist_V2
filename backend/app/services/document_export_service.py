"""
Document Export Service

Exports teaching scripts to PDF and Word formats.
"""

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF
from fpdf.enums import XPos, YPos

from app.core.config import settings
from app.services.script_service import PresentationScript, SlideScript

logger = logging.getLogger(__name__)

# Chinese font paths (Noto CJK fonts)
CHINESE_FONT_PATHS = [
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/google-noto-cjk/NotoSansCJK-Regular.ttc",
    # macOS
    "/System/Library/Fonts/PingFang.ttc",
    "/Library/Fonts/Arial Unicode.ttf",
    # Windows
    "C:/Windows/Fonts/msyh.ttc",
    "C:/Windows/Fonts/simsun.ttc",
]


def _find_chinese_font() -> str | None:
    """Find an available Chinese font on the system."""
    for font_path in CHINESE_FONT_PATHS:
        if Path(font_path).exists():
            return font_path
    return None


class ScriptPDF(FPDF):
    """Custom PDF class for script export with UTF-8 support."""

    def __init__(self, title: str):
        super().__init__()
        self.title = title
        self._chinese_font_loaded = False
        self._setup_fonts()

    def _setup_fonts(self):
        """Setup fonts for Chinese/UTF-8 support."""
        self.set_auto_page_break(auto=True, margin=15)

        # Try to load Chinese font
        font_path = _find_chinese_font()
        if font_path:
            try:
                # fpdf2 supports TTF/TTC fonts with Unicode
                self.add_font("NotoSansCJK", "", font_path)
                self.add_font("NotoSansCJK", "B", font_path)
                self.add_font("NotoSansCJK", "I", font_path)
                self._chinese_font_loaded = True
                logger.info(f"Chinese font loaded: {font_path}")
            except Exception as e:
                logger.warning(f"Failed to load Chinese font {font_path}: {e}")
                self._chinese_font_loaded = False
        else:
            logger.warning(
                "No Chinese font found, PDF may not display Chinese characters correctly"
            )

    def _get_font_name(self, style: str = "") -> str:
        """Get the appropriate font name based on availability."""
        if self._chinese_font_loaded:
            return "NotoSansCJK"
        return "Helvetica"

    def set_font_safe(self, style: str = "", size: int = 12):
        """Set font with fallback for Chinese support."""
        font_name = self._get_font_name(style)
        # Map style to fpdf style codes
        style_code = ""
        if "B" in style.upper():
            style_code += "B"
        if "I" in style.upper():
            style_code += "I"
        self.set_font(font_name, style_code, size)

    def header(self):
        """PDF header."""
        self.set_font_safe("B", 10)
        self.set_text_color(128, 128, 128)
        # Truncate title for header
        display_title = self.title[:50] if len(self.title) > 50 else self.title
        self.cell(0, 10, display_title, align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.set_draw_color(200, 200, 200)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(5)

    def footer(self):
        """PDF footer."""
        self.set_y(-15)
        self.set_font_safe("I", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


class DocumentExportService:
    """
    Service for exporting scripts to PDF and Word formats.

    Features:
    - PDF export with proper formatting and UTF-8/Chinese support
    - Word (.docx) export with styles
    - Section formatting for different script elements
    """

    def __init__(self, output_path: Path | None = None):
        """
        Initialize export service.

        Args:
            output_path: Base path for output files
        """
        self.output_path = output_path or settings.output_storage_path

    async def export_to_pdf(
        self,
        script: PresentationScript,
        presentation_id: str,
    ) -> Path:
        """
        Export script to PDF format.

        Args:
            script: PresentationScript to export
            presentation_id: Presentation ID for file naming

        Returns:
            Path to generated PDF file
        """
        logger.info(f"Exporting script to PDF: {script.title}")

        # Create output directory
        output_dir = self.output_path / presentation_id
        output_dir.mkdir(parents=True, exist_ok=True)
        output_file = output_dir / "script.pdf"

        # Create PDF
        pdf = ScriptPDF(script.title)
        pdf.alias_nb_pages()
        pdf.add_page()

        # Title page
        pdf.set_font_safe("B", 24)
        pdf.ln(20)  # Spacer
        pdf.multi_cell(0, 15, script.title, align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.set_font_safe("", 14)
        pdf.ln(10)
        pdf.cell(
            0,
            10,
            f"總時長：{script.total_minutes:.0f} 分鐘",
            align="C",
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )
        pdf.cell(
            0,
            10,
            f"投影片數量：{len(script.scripts)}",
            align="C",
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )

        pdf.set_font_safe("I", 10)
        pdf.ln(20)
        pdf.cell(
            0,
            10,
            f"生成時間：{script.generated_at.strftime('%Y-%m-%d %H:%M')}",
            align="C",
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )

        # Content pages
        for slide_script in script.scripts:
            self._add_slide_to_pdf(pdf, slide_script)

        # Save PDF
        pdf.output(str(output_file))

        logger.info(f"PDF exported: {output_file}")
        return output_file

    def _add_slide_to_pdf(self, pdf: ScriptPDF, slide: SlideScript):
        """Add a slide script to PDF."""
        pdf.add_page()

        # Slide header
        pdf.set_font_safe("B", 16)
        pdf.set_text_color(0, 102, 204)
        header = f"投影片 {slide.slide_index + 1}：{slide.slide_title}"
        pdf.multi_cell(0, 10, header, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.set_font_safe("I", 11)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(
            0,
            8,
            f"預估時間：{slide.estimated_minutes:.1f} 分鐘",
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )
        pdf.ln(5)

        # Lecture content
        pdf.set_font_safe("B", 12)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 8, "講課內容：", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        pdf.set_font_safe("", 11)
        content = slide.lecture_content or "(無內容)"
        self._add_multiline_text(pdf, content)
        pdf.ln(5)

        # Teaching tips
        if slide.teaching_tips:
            pdf.set_font_safe("B", 12)
            pdf.set_text_color(0, 128, 0)
            pdf.cell(0, 8, "教學提示：", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            pdf.set_font_safe("", 10)
            pdf.set_text_color(0, 0, 0)
            for tip in slide.teaching_tips:
                pdf.multi_cell(0, 6, f"  • {tip}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(3)

        # Interactive Q&A
        if slide.interaction_qa:
            pdf.set_font_safe("B", 12)
            pdf.set_text_color(204, 102, 0)
            pdf.cell(0, 8, "互動問答：", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            pdf.set_text_color(0, 0, 0)
            for qa in slide.interaction_qa:
                pdf.set_font_safe("B", 10)
                pdf.multi_cell(0, 6, f"  問：{qa.question}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

                if qa.expected_answers:
                    pdf.set_font_safe("I", 10)
                    pdf.multi_cell(
                        0,
                        6,
                        f"    預期答案：{', '.join(qa.expected_answers)}",
                        new_x=XPos.LMARGIN,
                        new_y=YPos.NEXT,
                    )
            pdf.ln(3)

        # Transition
        if slide.transition:
            pdf.set_font_safe("B", 12)
            pdf.set_text_color(128, 0, 128)
            pdf.cell(0, 8, "過場銜接：", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            pdf.set_font_safe("I", 10)
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(0, 6, slide.transition, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def _add_multiline_text(self, pdf: ScriptPDF, text: str, max_chars: int = 2000):
        """Add multiline text with proper UTF-8 handling."""
        # Truncate very long text
        if len(text) > max_chars:
            text = text[:max_chars] + "..."

        # With UTF-8 font support, no encoding conversion needed
        pdf.multi_cell(0, 6, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    async def export_to_docx(
        self,
        script: PresentationScript,
        presentation_id: str,
    ) -> Path:
        """
        Export script to Word (.docx) format.

        Args:
            script: PresentationScript to export
            presentation_id: Presentation ID for file naming

        Returns:
            Path to generated DOCX file
        """
        logger.info(f"Exporting script to DOCX: {script.title}")

        # Create output directory
        output_dir = self.output_path / presentation_id
        output_dir.mkdir(parents=True, exist_ok=True)
        output_file = output_dir / "script.docx"

        # Create document
        doc = Document()

        # Setup styles
        self._setup_docx_styles(doc)

        # Title
        title_para = doc.add_heading(script.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(
            f"總時長：{script.total_minutes:.0f} 分鐘 | 投影片數量：{len(script.scripts)}"
        )
        meta_run.font.size = Pt(12)
        meta_run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()  # Spacer

        # Table of contents hint
        toc_para = doc.add_paragraph()
        toc_para.add_run("目錄：").bold = True
        for slide in script.scripts:
            doc.add_paragraph(
                f"投影片 {slide.slide_index + 1}：{slide.slide_title} "
                f"({slide.estimated_minutes:.1f} 分鐘)",
                style="List Bullet",
            )

        doc.add_page_break()

        # Add each slide
        for slide_script in script.scripts:
            self._add_slide_to_docx(doc, slide_script)
            doc.add_page_break()

        # Save document
        doc.save(str(output_file))

        logger.info(f"DOCX exported: {output_file}")
        return output_file

    def _setup_docx_styles(self, doc: Document):
        """Setup custom styles for the document."""
        styles = doc.styles

        # Ensure Heading styles exist with proper formatting
        try:
            h1_style = styles["Heading 1"]
            h1_style.font.size = Pt(18)
            h1_style.font.color.rgb = RGBColor(0, 102, 204)
        except KeyError:
            pass

        try:
            h2_style = styles["Heading 2"]
            h2_style.font.size = Pt(14)
            h2_style.font.color.rgb = RGBColor(0, 128, 0)
        except KeyError:
            pass

    def _add_slide_to_docx(self, doc: Document, slide: SlideScript):
        """Add a slide script to Word document."""

        # Slide header
        header = f"投影片 {slide.slide_index + 1}：{slide.slide_title}"
        doc.add_heading(header, level=1)

        # Time estimate
        time_para = doc.add_paragraph()
        time_run = time_para.add_run(f"⏱️ 預估時間：{slide.estimated_minutes:.1f} 分鐘")
        time_run.font.italic = True
        time_run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()  # Spacer

        # Lecture content
        doc.add_heading("📖 講課內容", level=2)
        content_para = doc.add_paragraph(slide.lecture_content or "(無內容)")
        content_para.paragraph_format.line_spacing = 1.5

        # Teaching tips
        if slide.teaching_tips:
            doc.add_heading("💡 教學提示", level=2)
            for tip in slide.teaching_tips:
                doc.add_paragraph(tip, style="List Bullet")

        # Interactive Q&A
        if slide.interaction_qa:
            doc.add_heading("❓ 互動問答", level=2)
            for qa in slide.interaction_qa:
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f"問：{qa.question}")
                q_run.bold = True

                if qa.expected_answers:
                    a_para = doc.add_paragraph()
                    a_para.paragraph_format.left_indent = Inches(0.5)
                    a_run = a_para.add_run(f"預期答案：{', '.join(qa.expected_answers)}")
                    a_run.italic = True

        # Transition
        if slide.transition:
            doc.add_heading("🔗 過場銜接", level=2)
            trans_para = doc.add_paragraph(slide.transition)
            trans_para.paragraph_format.left_indent = Inches(0.25)
            for run in trans_para.runs:
                run.italic = True


# Global instance
_export_service: DocumentExportService | None = None


def get_export_service() -> DocumentExportService:
    """Get or create global export service instance."""
    global _export_service
    if _export_service is None:
        _export_service = DocumentExportService()
    return _export_service
